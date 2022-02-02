from docx import Document
from docx.shared import Inches
from docx.shared import Cm
import glob
from os import listdir
from os.path import isfile, join
from time import sleep
from tqdm import tqdm

import logging

logging.basicConfig(filename='app.log', filemode='w', format='%(name)s - %(levelname)s - %(message)s')
logging.warning('This will get logged to a file')


mypath = 'c:\\Users\\mohamad.alghish\\Desktop\\example\\out_new'
docx_path = "c:\\Users\\mohamad.alghish\\Desktop\\example\\out_new\\*.docx"
footer_png = 'C:\\Users\\mohamad.alghish\\Desktop\\Python\\pardocx\\footer.PNG'

# All files ending with .docx
fullpath = glob.glob(docx_path) 

with tqdm(total=len(fullpath)) as pbar:
	for f in fullpath:
		try:
			document = Document(f)
			section = document.sections[0]
			footer = section.footer
			footer.bottom_margin = Inches(0.0)
			footer.footer_distance = -Inches(1.0)
			paragraph = footer.paragraphs[0]
			paragraph.paragraph_format.left_indent = -Inches(0.8)
			run = paragraph.add_run()
			run.add_picture(footer_png,width=Inches(8.7), height=Inches(1.3))

			document.save(f)
		except Exception as e:
			logging.warning(f)

		pbar.update(1)
	



# onlyfiles = [f for f in listdir(mypath) if isfile(join(mypath, f))]
# # print(onlyfiles)



# # print(fullpath)

# for f in fullpath:

# 	document = Document(f)
# 	section = document.sections[0]
# 	footer = section.footer
# 	footer.bottom_margin = Inches(0.0)
# 	footer.footer_distance = -Inches(1.0)
# 	paragraph = footer.paragraphs[0]
# 	paragraph.paragraph_format.left_indent = -Inches(0.8)
# 	run = paragraph.add_run()
# 	run.add_picture(footer_png,width=Inches(8.7), height=Inches(1.3))

# 	document.save(f)




